import webbrowser
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx2pdf import convert
import os
# Provide the Folder path where you want to save the files
output = Path('H:\max') # Important
File_name = 'Test'
heading_text = "[ONLINE] UNLV vs Eastern Washington Live Stream NCAA Collage Football 02 September 2021 Full HD Coverage."
after_heading_text = "Welcome To Watch UNLV vs. Eastern Washington: How to live stream, TV channel, start time for Thursday's NCAA Football game. How to watch UNLV vs. Eastern Washington football game"
link_text = 'https://www.youtube.com/'

description_text = ''''Watch UNLV vs. Eastern Washington: How to live stream, TV channel, start time for Thursday's NCAA Football game
How to watch UNLV vs. Eastern Washington football game
Who's Playing
Eastern Washington @ UNLV

Last Season Records: UNLV 0-6; Eastern Washington 5-2

What to Know
The UNLV Rebels will play against a Division II opponent, the Eastern Washington Eagles, in an early-season tune-up Thursday at 10 p.m. ET at Allegiant Stadium. Coming off of an uninspired 0-6 last-season record, the Rebels have set their aspirations higher this year.

The neutral point spread forecasts a close one for these two. We'll see if the game is as close as the oddsmakers expect or if one of these teams has a surprise blowout in them.

How To Watch
When: Thursday at 10 p.m. ET
Where: Allegiant Stadium -- Paradise, Nevada
TV: STADIUM
Follow: CBS Sports App
Ticket Cost: $15.00
Odds
The Rebels are a slight 2-point favorite against the Eagles, according to the latest college football odds.

The oddsmakers had a good feel for the line for this one, as the game opened with the Rebels as a 3-point favorite.

Over/Under: -110

See college football picks for every single game, including this one, from SportsLine's advanced computer model. Get picks now.

Series History
This is the first time these teams have played each other within the last six years.

How to Watch Eastern Washington Eagles  at UNLV Rebels: Live Stream, TV Channel, Start Time
UNLV looks to pick up its first win in two years when it hosts Eastern Washington in the season opener Thursday night.

UNLV made a splash by signing transfer Tate Martell right before school started, but in a mild surprise, he was nowhere to be found when the Rebels released their depth chart. The Rebels, who have not won a game since 2019, need a spark and many thought he would be it. Instead, they will wait for him to play quarterback at least another week.

How to Watch:

Date: Sept. 2, 2021

Time: 10:00 p.m. ET

TV: Stadium

You can stream the game on fuboTV: Start with a 7-day free trial!

UNLV has a rough start to its schedule, and if it can't find a way to beat Eastern Washington, it could be a while before the Rebels win. After their opening game, they travel to No. 25 Arizona State before hosting Iowa State. They are going to be huge underdogs in those games, so they need to find a way to beat the Eagles.

Eastern Washington went 5-2 in the spring of 2021 after the fall schedule was postponed due to COVID-19. The Eagles lost to North Dakota State in the first round of the playoffs after losing just one game in the regular season.

The Eagles are not going to be an easy game for the Rebels, and if you are looking for an FCS team to upset an FBS team this would be the game to watch. The Rebels haven't had a winning season since 2013 and have had a hard time winning even four games in a year. Eastern Washington will give them everything they can handle and could easily win this game.
'''

output_directory = Path(output)
output_directory_pdf = output / 'pdf'
output_directory_docx = output / 'docx'


Path(output_directory).mkdir(parents=True, exist_ok=True)
Path(output_directory_docx).mkdir(parents=True, exist_ok=True)
Path(output_directory_pdf).mkdir(parents=True, exist_ok=True)


docx_file_location = Path(output_directory_docx, File_name+'.docx')
pdf_file_location = Path(output_directory_pdf, File_name+'.pdf')

doc = Document()


def make_style(doc, style_name, font_size, font_name):
    font_styles = doc.styles
    font_charstyle = font_styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(font_size)
    font_object.name = font_name


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


make_style(doc, 'headingStyle', 26, 'Cambria')
make_style(doc, 'afterheadingStyle', 12, 'Calibri')
heading = doc.add_paragraph().add_run(heading_text, style='headingStyle').bold = True
after_heading = doc.add_paragraph().add_run(after_heading_text, style='afterheadingStyle').bold = True
link = doc.add_paragraph().add_run(link_text, style='afterheadingStyle').bold = True
description = doc.add_paragraph().add_run(description_text, style='afterheadingStyle').bold = True

from docx.enum.text import WD_ALIGN_PARAGRAPH

# heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 0 for left, 1 for center, 2 for right
# heading.alignment = 1
# after_heading.alignment = 1
# link.alignment = 1
# description.alignment = 1


doc.save(docx_file_location)



convert(docx_file_location, pdf_file_location)
print(f"Files successfully generated in '{output_directory}'")

webbrowser.open(output_directory)  # Opens the output folder.
